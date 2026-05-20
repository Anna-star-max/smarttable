from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from typing import Optional
import os
import uuid
import json
from datetime import datetime

from backend.models.schemas import AnalyzeRequest, AnalyzeResponse, HistoryRecord
from backend.services.minimax import minimax_service
from backend.services.excel_generator import excel_generator

def fix_date_to_current_year(date_str):
    if not date_str:
        return None
    current_year = datetime.now().year
    formats = ["%Y-%m-%d", "%Y/%m/%d", "%Y年%m月%d日", "%m月%d日"]
    for fmt in formats:
        try:
            parsed = datetime.strptime(date_str, fmt)
            if fmt == "%m月%d日":
                return parsed.replace(year=current_year)
            if parsed.year < current_year:
                return parsed.replace(year=current_year)
            return parsed
        except:
            pass
    return None

router = APIRouter()

# 歷史記錄儲存（記憶體，實際環境可用資料庫）
history_file = "./output/history.json"

def load_history():
    if os.path.exists(history_file):
        with open(history_file, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_history(records):
    os.makedirs("./output", exist_ok=True)
    with open(history_file, "w", encoding="utf-8") as f:
        json.dump(records, f, ensure_ascii=False, indent=2, default=str)

@router.post("/api/analyze", response_model=AnalyzeResponse)
async def analyze_text(request: AnalyzeRequest):
    try:
        import os
        api_key = os.getenv("MINIMAX_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="MINIMAX_API_KEY not loaded from .env")

        result = await minimax_service.analyze_text(request.text)

        content_type = result.get("content_type", "unknown")

        # 通用處理：修正日期
        meeting_overview = result.get("meeting_overview") or {}
        date_val = meeting_overview.get("日期")
        if date_val:
            fixed = fix_date_to_current_year(date_val)
            meeting_overview["日期"] = fixed.strftime("%m月%d日") if fixed else datetime.now().strftime("%m月%d日")
        else:
            meeting_overview["日期"] = datetime.now().strftime("%m月%d日")

        for item in result.get("action_items", []):
            if item.get("截止日期"):
                fixed = fix_date_to_current_year(item["截止日期"])
                item["截止日期"] = fixed.strftime("%m月%d日") if fixed else "-"

        # 根據 content_type 決定 Excel 標題與生成方式
        if content_type == "data_table":
            meeting_topic = result.get("data_table", {}).get("title") or "數據表格"
        elif content_type == "contract":
            meeting_topic = result.get("contract_info", {}).get("標題") or "合同分析"
        elif content_type == "report":
            meeting_topic = result.get("report_info", {}).get("標題") or "報告分析"
        elif content_type == "lecture":
            meeting_topic = result.get("lecture_info", {}).get("標題") or "課程摘要"
        elif content_type == "financial":
            meeting_topic = result.get("financial_info", {}).get("標題") or "財務分析"
        else:
            meeting_topic = (result.get("meeting_overview") or {}).get("主題") or "未命名會議"

        excel_filename = excel_generator.generate(result, meeting_topic)

        record = {
            "id": uuid.uuid4().hex,
            "created_at": datetime.now().isoformat(),
            "meeting_topic": meeting_topic,
            "excel_filename": excel_filename,
            "data": result
        }

        history = load_history()
        history.insert(0, record)
        save_history(history[:100])

        return AnalyzeResponse(
            success=True,
            data=result,
            excel_filename=excel_filename,
            created_at=datetime.fromisoformat(record["created_at"])
        )
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        raise HTTPException(status_code=500, detail=str(e) + "\n" + traceback.format_exc())

@router.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    ext = file.filename.split('.')[-1].lower()
    if ext == "txt":
        content = await file.read()
        text = content.decode("utf-8")
    elif ext == "docx":
        from docx import Document
        import io
        content = await file.read()
        doc = Document(io.BytesIO(content))
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        text = "\n".join(text_parts)
    elif ext == "pdf":
        import pdfplumber
        import io
        content = await file.read()
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = "\n".join(text_parts)
    elif ext in ["xlsx", "xls"]:
        import pandas as pd
        import io
        content = await file.read()
        # 读取所有工作表
        xls = pd.ExcelFile(io.BytesIO(content))
        text_parts = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            text_parts.append(f"【{sheet_name}】\n{df.to_string(index=False)}")
        text = "\n\n".join(text_parts)
    else:
        raise HTTPException(status_code=400, detail="僅支援 .txt、.docx、.pdf、.xlsx、.xls、.png、.jpg、.jpeg 檔案")

    result = await minimax_service.analyze_text(text)

    content_type = result.get("content_type", "unknown")

    # 通用處理：修正日期
    meeting_overview = result.get("meeting_overview") or {}
    date_val = meeting_overview.get("日期")
    if date_val:
        fixed = fix_date_to_current_year(date_val)
        meeting_overview["日期"] = fixed.strftime("%m月%d日") if fixed else datetime.now().strftime("%m月%d日")
    else:
        meeting_overview["日期"] = datetime.now().strftime("%m月%d日")

    for item in result.get("action_items", []):
        if item.get("截止日期"):
            fixed = fix_date_to_current_year(item["截止日期"])
            item["截止日期"] = fixed.strftime("%m月%d日") if fixed else "-"

    # 根據 content_type 決定 Excel 標題與生成方式
    if content_type == "bug_report":
        meeting_topic = result.get("bug_report", {}).get("標題") or "Bug問題分析"
    elif content_type == "data_table":
        meeting_topic = result.get("data_table", {}).get("title") or "數據表格"
    elif content_type == "data_table":
        meeting_topic = result.get("data_table", {}).get("title") or "數據表格"
    elif content_type == "contract":
        meeting_topic = result.get("contract_info", {}).get("標題") or "合同分析"
    elif content_type == "report":
        meeting_topic = result.get("report_info", {}).get("標題") or "報告分析"
    elif content_type == "lecture":
        meeting_topic = result.get("lecture_info", {}).get("標題") or "課程摘要"
    elif content_type == "financial":
        meeting_topic = result.get("financial_info", {}).get("標題") or "財務分析"
    else:
        meeting_topic = (result.get("meeting_overview") or {}).get("主題") or file.filename

    excel_filename = excel_generator.generate(result, meeting_topic)

    record = {
        "id": uuid.uuid4().hex,
        "created_at": datetime.now().isoformat(),
        "meeting_topic": meeting_topic,
        "excel_filename": excel_filename,
        "data": result
    }

    history = load_history()
    history.insert(0, record)
    save_history(history[:100])

    return AnalyzeResponse(
        success=True,
        data=result,
        excel_filename=excel_filename,
        created_at=datetime.fromisoformat(record["created_at"])
    )

@router.post("/api/regenerate")
async def regenerate_excel(data: dict, meeting_topic: str = "未命名會議"):
    """
    Regenerate Excel from modified data without AI analysis.
    """
    try:
        if not data:
            raise HTTPException(status_code=400, detail="Missing data")

        excel_filename = excel_generator.generate(data, meeting_topic)
        return {"excel_filename": excel_filename}
    except Exception as e:
        import traceback
        raise HTTPException(status_code=500, detail=str(e) + "\n" + traceback.format_exc())

@router.get("/api/download/{filename}")
async def download_file(filename: str):
    filepath = os.path.join("./output", filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="檔案不存在")
    return FileResponse(filepath, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", filename=filename)

@router.get("/api/history")
async def get_history():
    history = load_history()
    return {"records": history}

@router.get("/api/records/{record_id}")
async def get_record(record_id: str):
    history = load_history()
    for record in history:
        if record["id"] == record_id:
            return record
    raise HTTPException(status_code=404, detail="記錄不存在")

@router.delete("/api/records/{record_id}")
async def delete_record(record_id: str):
    history = load_history()
    original_length = len(history)
    history = [r for r in history if r["id"] != record_id]

    if len(history) == original_length:
        raise HTTPException(status_code=404, detail="記錄不存在")

    # 刪除關聯的 Excel 檔案
    for r in load_history():
        if r["id"] == record_id:
            excel_path = os.path.join("./output", r.get("excel_filename", ""))
            if os.path.exists(excel_path):
                os.remove(excel_path)
            break

    save_history(history)
    return {"success": True, "message": "已刪除"}